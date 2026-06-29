# -*- coding: utf-8 -*-
"""Generate formatted Word analysis report from computed financial data.

Parameterized engine — takes PL/BS/CF results + company info →
produces a professional analysis report with risk assessment and recommendations.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime


def generate_analysis_report(company_name, period_label, bs_date, pl, bs, cf,
                              opening_bs, output_path):
    """Generate a formatted Word analysis report.

    Args:
        company_name: e.g. '来宾市顺恒废旧汽车回收有限公司'
        period_label: e.g. '2026年1-5月'
        bs_date: e.g. '2026年5月31日'
        pl: PL result dict from CalcEngine.calculate_pl()
        bs: BS result dict from CalcEngine.calculate_bs()
        cf: CF result dict from CalcEngine.calculate_cf()
        opening_bs: dict of opening balance sheet values
        output_path: where to save the .docx file
    """
    doc = Document()

    # ---- Page setup ----
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5

    # ---- Helpers ----
    def _v(val, default=0.0):
        """Safely get float value."""
        if val is None:
            return default
        return float(val)

    def _fmt(val):
        """Format a number with commas and 2 decimal places."""
        if val is None:
            return '—'
        return f'{float(val):,.2f}'

    def add_title(text, level=0):
        if level == 0:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(text)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(18)
            run.bold = True
        else:
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        return doc.paragraphs[-1]

    def add_para(text, bold=False, indent=True, size=10.5):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        if indent:
            p.paragraph_format.first_line_indent = Pt(21)
        run = p.add_run(text)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(size)
        run.bold = bold
        return p

    def add_table(headers, rows, col_widths=None):
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, hdr in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(hdr))
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(9)
            run.bold = True
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
            cell._tc.get_or_add_tcPr().append(shading)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                if ci == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci <= 2 else WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(str(val) if val is not None else '')
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(9)
        if col_widths:
            for row in table.rows:
                for ci, w in enumerate(col_widths):
                    if ci < len(row.cells):
                        row.cells[ci].width = Cm(w)
        doc.add_paragraph()
        return table

    def add_warning(text):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(21)
        run = p.add_run(text)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)
        run.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        return p

    # ================================================================
    # Compute key metrics for dynamic analysis
    # ================================================================
    rev = _v(pl.get('revenue'))
    cogs = _v(pl.get('cogs'))
    gross = rev - cogs
    gross_margin = (gross / rev * 100) if rev > 0 else 0
    net_profit = _v(pl.get('net_profit'))
    net_margin = (net_profit / rev * 100) if rev > 0 else 0
    admin_exp = _v(pl.get('admin_exp'))
    fin_exp = _v(pl.get('fin_exp'))
    surcharges = _v(pl.get('surcharges'))

    total_assets = _v(bs.get('total_assets'))
    total_liab = _v(bs.get('total_liab'))
    total_equity = _v(bs.get('total_equity'))
    asset_liability_ratio = (total_liab / total_assets * 100) if total_assets > 0 else 0
    cash = _v(bs.get('cash'))
    other_pay = _v(bs.get('other_pay'))
    is_insolvent = total_equity < 0

    op_cf = _v(cf.get('经营活动产生的现金流量净额'))
    cash_change = _v(cf.get('现金及现金等价物净增加额'))
    depr = _v(pl.get('_detail', {}).get('depreciation', 0))

    # Detect anomalies
    detail = pl.get('_detail', {})
    annot = _fmt  # alias for formatting

    # ================================================================
    # CONTENT
    # ================================================================

    add_title(company_name)
    add_title(f'{period_label}财务报表分析说明')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(f'编制日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)

    # ---- Section 1: Basis ----
    add_title('一、报表编制基础', level=1)
    add_para(f'本报告基于{company_name}{period_label}的财务数据编制，执行《小企业会计准则》。')
    add_para('数据来源包括：销售发票、成本费用发票、银行流水及工资薪金表。报表生成工具自动抽取、映射、计算后产出三大报表及本分析报告。')

    # ---- Section 2: Operating Results ----
    add_title('二、整体经营情况', level=1)

    add_table(
        ['指标', f'{period_label}累计（元）', '说明'],
        [
            ['营业收入', _fmt(rev), f'毛利率 {gross_margin:.1f}%'],
            ['营业成本', _fmt(cogs), f'成本率 {cogs/rev*100:.1f}%' if rev > 0 else '—'],
            ['毛利', _fmt(gross), f'毛利率 {gross_margin:.1f}%'],
            ['税金及附加', _fmt(surcharges), '增值税附加税'],
            ['管理费用', _fmt(admin_exp), f'占收入 {admin_exp/rev*100:.1f}%' if rev > 0 else '—'],
            ['财务费用', _fmt(fin_exp), '银行账户管理费及利息'],
            ['净利润', _fmt(net_profit), f'净利率 {net_margin:.1f}%'],
        ],
        col_widths=[3.5, 3.5, 7.0]
    )

    # Dynamic warnings based on metrics
    if net_profit < 0:
        add_title('⚠ 本期亏损', level=2)
        add_warning(f'{period_label}净利润为{_fmt(net_profit)}元，公司处于亏损状态。')
        add_para('亏损可能原因包括：收入不足、成本过高、管理费用膨胀或一次性支出。建议逐月对比收入成本趋势，找出亏损根因。')
    elif net_margin < 5:
        add_title('⚠ 盈利偏薄', level=2)
        add_warning(f'净利率仅{net_margin:.1f}%，虽然盈利但利润空间很小。')
        add_para('微利经营对市场波动和成本上涨非常敏感，建议通过提高销售价格、降低采购成本或控制费用来提升利润率。')

    if admin_exp > rev * 0.3 and rev > 0:
        add_title('⚠ 管理费用占比偏高', level=2)
        add_warning(f'管理费用占营业收入的{admin_exp/rev*100:.1f}%，超出合理范围。')
        add_para('管理费用包含工资、社保、折旧、办公费等。建议核查各项明细，确认是否有异常大额支出或费用跨期错配。')

    # ---- Section 3: Balance Sheet ----
    add_title('三、资产负债状况', level=1)

    bs_rows = [
        ['货币资金', _fmt(bs.get('cash')), _fmt(opening_bs.get('货币资金'))],
        ['应收账款', _fmt(bs.get('ar')), _fmt(opening_bs.get('应收账款'))],
        ['预付账款', _fmt(bs.get('prepay')), _fmt(opening_bs.get('预付账款'))],
        ['其他应收款', _fmt(bs.get('other_recv')), _fmt(opening_bs.get('其他应收款'))],
        ['存货', _fmt(bs.get('inventory')), _fmt(opening_bs.get('存货'))],
        ['流动资产合计', _fmt(bs.get('curr_assets')), _fmt(opening_bs.get('流动资产合计'))],
        ['固定资产净值', _fmt(bs.get('fa_net')), _fmt(opening_bs.get('固定资产账面价值'))],
        ['长期待摊费用', _fmt(bs.get('ltd')), _fmt(opening_bs.get('长期待摊费用'))],
        ['资产总计', _fmt(total_assets), _fmt(opening_bs.get('资产总计'))],
        ['', '', ''],
        ['应付账款', _fmt(bs.get('ap')), _fmt(opening_bs.get('应付账款'))],
        ['预收账款', _fmt(bs.get('pr')), _fmt(opening_bs.get('预收账款'))],
        ['应付职工薪酬', _fmt(bs.get('payroll_payable')), _fmt(opening_bs.get('应付职工薪酬'))],
        ['应交税费', _fmt(bs.get('tax_payable')), _fmt(opening_bs.get('应交税费'))],
        ['其他应付款', _fmt(other_pay), _fmt(opening_bs.get('其他应付款'))],
        ['负债合计', _fmt(total_liab), _fmt(opening_bs.get('负债合计'))],
        ['', '', ''],
        ['实收资本', _fmt(bs.get('capital')), _fmt(opening_bs.get('实收资本（或股本）'))],
        ['未分配利润', _fmt(bs.get('undist_profit_end')), _fmt(bs.get('undist_profit_beg'))],
        ['所有者权益合计', _fmt(total_equity), _fmt(opening_bs.get('所有者权益（或股东权益）合计'))],
    ]
    add_table(
        ['项目', '期末余额（元）', '年初余额（元）'],
        bs_rows,
        col_widths=[4.0, 3.5, 3.5]
    )

    # Dynamic BS warnings
    risks = []  # Collect risks for summary section

    if is_insolvent:
        add_title('⚠ 资不抵债——持续经营重大疑虑', level=2)
        add_warning(f'所有者权益为{_fmt(total_equity)}元，公司已处于资不抵债状态。资产负债率高达{asset_liability_ratio:.1f}%。')
        add_para('从法律角度看，即使将所有资产变卖变现，也不足以偿还全部债务。虽然小企业通常依靠股东或关联方持续资金支持维持运营，但这种状态已构成持续经营的重大不确定性。')
        risks.append(('资不抵债', '严重',
                      f'所有者权益{_fmt(total_equity)}，资产负债率{asset_liability_ratio:.1f}%，持续经营存在重大不确定性'))

    if other_pay > total_assets * 0.5 and other_pay > 0:
        add_title('⚠ 其他应付款占比过高——关联方借款是主要资金来源', level=2)
        ratio = other_pay / total_liab * 100 if total_liab > 0 else 0
        add_warning(f'其他应付款{_fmt(other_pay)}元，占负债总额的{ratio:.1f}%，公司实质上依赖股东/关联方借款维持运营。')
        add_para(f'现金余额仅{_fmt(cash)}元。如果关联方要求大额还款，公司将面临严重的流动性危机。建议与关联方就借款金额、利率、还款期限等签订书面协议。')
        risks.append(('关联方依赖', '严重',
                      f'其他应付款{_fmt(other_pay)}占负债{ratio:.1f}%，完全依赖关联方资金输血'))

    # Check for stagnant AR/AP/Inventory
    ar = _v(bs.get('ar'))
    ar_open = _v(opening_bs.get('应收账款'))
    inv = _v(bs.get('inventory'))
    inv_open = _v(opening_bs.get('存货'))
    prepay = _v(bs.get('prepay'))
    prepay_open = _v(opening_bs.get('预付账款'))

    if abs(ar - ar_open) < 0.5 and abs(inv - inv_open) < 0.5 and abs(prepay - prepay_open) < 0.5 \
       and (ar > 0 or inv > 0 or prepay > 0):
        add_title('⚠ 应收、预付、存货余额未变动', level=2)
        add_warning(f'应收账款{_fmt(ar)}、预付账款{_fmt(prepay)}、存货{_fmt(inv)}——期末余额与年初余额基本一致，整个期间没有任何变动。')
        add_para('可能原因：未做往来重分类和账龄分析；未对存货进行实地盘点；应收账款账龄可能较长，存在坏账风险但未计提减值准备。建议尽快对往来款项和存货进行清理核实。')
        risks.append(('往来款长期未动', '中等',
                      '应收/预付/存货余额与年初一致，可能存在账龄长、未对账、坏账等问题'))

    # Cash flow check
    if _v(cash) < _v(total_liab) * 0.1 and _v(cash) < 500000:
        risks.append(('现金流断裂风险', '较高' if _v(cash) < 100000 else '中等',
                      f'现金仅{_fmt(cash)}元，短期偿债能力不足'))

    # ---- Section 4: Cash Flow ----
    add_title('四、现金流量分析', level=1)

    add_table(
        ['项目', '金额（元）'],
        [
            ['一、经营活动现金流量净额', _fmt(cf.get('经营活动产生的现金流量净额'))],
            ['二、投资活动现金流量净额', _fmt(cf.get('投资活动产生的现金流量净额'))],
            ['三、筹资活动现金流量净额', _fmt(cf.get('筹资活动产生的现金流量净额'))],
            ['四、现金净增加额', _fmt(cf.get('现金及现金等价物净增加额'))],
            ['加：期初现金余额', _fmt(cf.get('加：期初现金余额'))],
            ['五、期末现金余额', _fmt(cf.get('五、期末现金余额'))],
        ],
        col_widths=[8.0, 6.0]
    )

    doc.add_paragraph()

    # Indirect method reconciliation
    add_table(
        ['间接法：净利润→经营现金流', '金额（元）'],
        [
            ['净利润', _fmt(net_profit)],
            ['加：固定资产折旧', _fmt(depr)],
            ['加：长期待摊费用摊销', _fmt(_v(pl.get('_detail', {}).get('ltd_amort', 0)))],
            ['减：经营性应收项目增加', _fmt(-_v(cf.get('经营性应收项目的减少（减：增加）')))],
            ['加：经营性应付项目增加', _fmt(_v(cf.get('经营性应付项目的增加（减：减少）')))],
            ['= 经营活动现金流量净额', _fmt(op_cf)],
        ],
        col_widths=[8.0, 6.0]
    )

    # Cash flow quality assessment
    cf_quality = op_cf / net_profit if net_profit != 0 else None
    if cf_quality is not None and cf_quality < 0.5:
        add_title('⚠ 经营活动现金流偏弱，盈利质量堪忧', level=2)
        add_warning(f'经营活动现金净流入仅{_fmt(op_cf)}元，净利润{_fmt(net_profit)}元，现金流仅为利润的{cf_quality*100:.0f}%。')
        add_para('间接法揭示了原因：公司虽有利润，但营运资金变动吞噬了大部分现金。换句话说，"账面上赚钱了"但钱未能形成可用的现金积累。这种盈利质量对短期偿债能力和日常运营资金保障构成严重威胁。')
        risks.append(('盈利质量差', '较高',
                      f'经营现金流仅{_fmt(op_cf)}，仅为净利润的{cf_quality*100:.0f}%'))

    # ---- Section 5: Tax ----
    add_title('五、税务分析', level=1)
    vat = detail.get('vat', 0)
    cj_tax = vat * 0.07
    jy_tax = vat * 0.03
    dj_tax = vat * 0.02

    add_table(
        ['税种', '计税依据', '金额（元）'],
        [
            ['增值税', '销项税-进项税', _fmt(vat)],
            ['城市维护建设税（7%）', '应纳增值税额', _fmt(cj_tax)],
            ['教育费附加（3%）', '应纳增值税额', _fmt(jy_tax)],
            ['地方教育附加（2%）', '应纳增值税额', _fmt(dj_tax)],
            ['税金及附加合计', '', _fmt(surcharges)],
            ['企业所得税', '应纳税所得额',
             '0.00（累计亏损，无需缴纳）' if _v(bs.get('undist_profit_end')) < 0 else _fmt(net_profit * 0.025)],
        ],
        col_widths=[5.0, 4.5, 5.0]
    )

    # ---- Section 6: Comprehensive Assessment ----
    add_title('六、综合评估', level=1)

    if risks:
        add_title('🔴 风险信号', level=2)
        add_table(
            ['风险', '严重程度', '详细说明'],
            [[r[0], r[1], r[2]] for r in risks],
            col_widths=[3.5, 2.0, 9.0]
        )

    # Positive signals
    positives = []
    if gross_margin > 30:
        positives.append(['毛利率良好', f'毛利率{gross_margin:.1f}%，主营业务具有一定的盈利能力'])
    if net_profit > 0:
        positives.append(['实现盈利', f'{period_label}累计净利润{_fmt(net_profit)}元，企业仍有经营价值'])
    if is_insolvent and total_equity > _v(opening_bs.get('所有者权益（或股东权益）合计', -1e10)):
        positives.append(['净资产缺口收窄', f'所有者权益在改善中，趋势向好'])
    if _v(bs.get('fa_net')) > 1000000:
        positives.append(['固定资产基础', f'固定资产净值{_fmt(bs.get("fa_net"))}元，具备一定的资产基础'])

    if positives:
        add_title('🟢 积极信号', level=2)
        add_table(
            ['方面', '说明'],
            positives,
            col_widths=[4.0, 10.5]
        )

    # ---- Section 7: Recommendations ----
    add_title('七、建议', level=1)

    recs = []
    if is_insolvent:
        recs.append(('与关联方明确债务安排',
                     f'{_fmt(other_pay)}元其他应付款应补签书面借款协议，明确金额、利率、还款计划。必要时将部分借款转为增资以改善资本结构。'))
        recs.append(('考虑股东增资',
                     f'实收资本仅{_fmt(bs.get("capital"))}元，累计亏损较大，增资是解决资不抵债问题的最直接方式。'))
    if cf_quality is not None and cf_quality < 0.5:
        recs.append(('编制现金流预算',
                     f'现金余额仅{_fmt(cash)}元，建议编制未来3-6个月的月度现金流预算，提前识别资金缺口。'))
    if abs(ar - ar_open) < 0.5 and ar > 0:
        recs.append(('加强往来款管理', '对应收账款进行账龄分析催收；对预付账款核实业务进度；对存货实地盘点确保账实相符。'))
    recs.append(('规范财务核算', '建议按月进行往来重分类、折旧计提、税费计算，避免期末集中处理；建立固定资产卡片和存货进销存台账。'))

    for i, (title, desc) in enumerate(recs, 1):
        add_para(f'{i}. {title}', bold=True)
        add_para(desc)

    # ---- Footer ----
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'— 本报告基于{period_label}账面数据自动生成，未经审计 —')
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ---- Save ----
    doc.save(output_path)
    return output_path


# ================================================================
# Standalone test (backward compatible)
# ================================================================
if __name__ == '__main__':
    print("This module is now parameterized. Use generate_analysis_report() instead.")
    print("See app.py '分析导出' page for integration.")
